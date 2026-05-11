"""
scripts/export_docx.py - Export final_report.md and advanced_strategy_analysis.md to DOCX.
Usage: python scripts/export_docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

REPO = Path(__file__).parent.parent


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_bold_run(para, text):
    run = para.add_run(text)
    run.bold = True
    return run


def parse_inline(para, text):
    """Handle **bold**, `code`, and plain text within a line."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            para.add_run(part)


def md_to_docx(md_path: Path, doc: Document, img_dir: Path | None = None):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip image lines (will note them as caption)
        if line.startswith("!["):
            alt = re.search(r"!\[([^\]]*)\]", line)
            caption = alt.group(1) if alt else "Chart"
            # Try to embed the image
            img_match = re.search(r"\(([^)]+\.png)\)", line)
            if img_match and img_dir:
                img_path = img_dir / img_match.group(1)
                if img_path.exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap = doc.add_paragraph(caption)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.runs[0].italic = True
                        cap.runs[0].font.size = Pt(9)
                    except Exception:
                        doc.add_paragraph(f"[Chart: {caption}]").italic = True
                else:
                    doc.add_paragraph(f"[Chart: {caption}]").italic = True
            i += 1
            continue

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:], level=1)

        # H2
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)

        # H3
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)

        # Horizontal rule
        elif line.strip() == "---":
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # Table (collect all rows)
        elif line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].split("|")[1:-1]]
                if not all(re.match(r"^[-:]+$", c) for c in row):
                    rows.append(row)
                i += 1
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        # Strip bold markers for table cells
                        clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
                        clean = re.sub(r"`([^`]+)`", r"\1", clean)
                        if ri == 0:
                            run = para.add_run(clean)
                            run.bold = True
                            set_cell_bg(cell, "D9E1F2")
                        else:
                            para.add_run(clean)
                        para.paragraph_format.space_after = Pt(2)
                        para.paragraph_format.space_before = Pt(2)
                doc.add_paragraph()
            continue

        # Blockquote
        elif line.startswith("> "):
            para = doc.add_paragraph(style="Quote")
            parse_inline(para, line[2:])

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            parse_inline(para, line[2:])

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            para = doc.add_paragraph(style="List Number")
            parse_inline(para, re.sub(r"^\d+\. ", "", line))

        # Italic line (starts and ends with *)
        elif line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(line.strip("*"))
            run.italic = True

        # Empty line
        elif line.strip() == "":
            pass

        # Normal paragraph
        else:
            para = doc.add_paragraph()
            parse_inline(para, line)

        i += 1


def make_docx(md_path: Path, out_path: Path, img_dir: Path | None = None):
    doc = Document()

    # Page margins
    from docx.shared import Cm

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Styles
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    md_to_docx(md_path, doc, img_dir=img_dir)

    doc.save(str(out_path))
    print(f"Saved: {out_path.name}")


if __name__ == "__main__":
    reports = REPO / "reports"

    make_docx(
        reports / "final_report.md",
        reports / "final_report.docx",
        img_dir=reports,
    )

    make_docx(
        reports / "advanced_strategy_analysis.md",
        reports / "advanced_strategy_analysis.docx",
        img_dir=reports,
    )
